from docx import Document
from docx.shared import Pt

document = Document()

#输出&保存
document.add_paragraph('测试一下')
document.save('app.docx')

#在一个段落中增加文字
paragraph = document.add_paragraph('增加，')

#增加文字
paragraph.add_run('增加文字')

#设置word字体大小
style = document.styles['Normal']
font = style.font
font.size = Pt(10)
document.save('app1.docx')